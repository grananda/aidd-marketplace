#!/usr/bin/env python3
"""Los scripts que producen un entregable **se ejecutan**, no solo compilan.

`py_compile` no ve un `NameError`: el fichero compila y revienta al usarlo. Paso
exactamente por eso --una variable mal escrita en `compute_kpis.py` llego a una
release y dejaba `aiba metrics` sin poder ejecutarse--, y ninguna de las nueve
comprobaciones lo detecto porque ninguna **llamaba** al script.

Esta lo hace: monta un proyecto minimo y lo ejecuta de punta a punta. No valida
las cifras --eso es de cada script-- sino que termine con exito y devuelva algo
parseable. Es un humo, y es el que faltaba.
"""
from __future__ import annotations

import json
import subprocess
import sys
import tempfile
from pathlib import Path

ROOT = Path(__file__).resolve().parents[2]
KPIS = ROOT / "plugins/aiba/skills/aiba-metrics/scripts/compute_kpis.py"
AUDIT = ROOT / "plugins/aisdd/skills/aisdd-specs/scripts/audit.py"
DF = ROOT / "plugins/aiba/skills/aiba-functional-design/scripts/gen_df_docx.py"

ACTIVIDAD = """# Registro de actividad AIDD

- 2026-09-03T10:00:00Z | user:dev | skill:aisdd-specs | ctx:HU-01 | run | note:-
- 2026-09-03T10:01:00Z | user:dev | skill:aisdd-specs | ctx:HU-01 | file:src/a.ts | note:-
- 2026-09-03T10:05:00Z | user:dev | skill:aisdd-specs | ctx:HU-01 | turn | note:dur=300s skills=1 files=1
"""

errors: list[str] = []


def proyecto(d: Path) -> None:
    (d / "docs").mkdir()
    (d / "openspec" / "audit" / "2026-09").mkdir(parents=True)
    (d / "openspec" / "changes" / "archive" / "viejo").mkdir(parents=True)
    (d / "docs" / "aidd-activity.md").write_text(ACTIVIDAD, encoding="utf-8")
    entradas = [
        {"command": "aisdd open change", "change_id": "nuevo", "id": "a",
         "timestamp": "2026-09-03T10:00:00Z", "corrects_archived": "viejo"},
        {"command": "aisdd close change", "change_id": "nuevo", "id": "b",
         "timestamp": "2026-09-03T12:00:00Z"},
    ]
    (d / "openspec" / "audit" / "2026-09" / "dev.jsonl").write_text(
        "\n".join(json.dumps(e) for e in entradas) + "\n", encoding="utf-8")


with tempfile.TemporaryDirectory() as tmp:
    d = Path(tmp)
    proyecto(d)

    # 1. `aiba metrics`: el informe completo, que es donde estaba el fallo.
    r = subprocess.run([sys.executable, str(KPIS), "--audit", "openspec/audit",
                        "--no-git", "--format", "json"],
                       cwd=d, capture_output=True, text=True, timeout=120)
    if r.returncode != 0:
        errors.append(f"compute_kpis.py falla al ejecutarse: {r.stderr.strip()[-400:]}")
    else:
        try:
            datos = json.loads(r.stdout)
        except json.JSONDecodeError as exc:
            errors.append(f"compute_kpis.py no devuelve JSON valido: {exc}")
        else:
            for clave in ("attended", "audit"):
                if clave not in datos:
                    errors.append(f"compute_kpis.py: falta la seccion '{clave}' en la salida")
            if datos.get("audit", {}).get("rework_total") != 1:
                errors.append("compute_kpis.py: no cuenta el retrabajo "
                              "(`corrects_archived`) que trae la auditoria de prueba")

    # 2. Y el formato humano, que recorre otro camino del codigo.
    r = subprocess.run([sys.executable, str(KPIS), "--audit", "openspec/audit",
                        "--no-git", "--format", "md"],
                       cwd=d, capture_output=True, text=True, timeout=120)
    if r.returncode != 0:
        errors.append(f"compute_kpis.py --format md falla: {r.stderr.strip()[-400:]}")

    # 3. `audit.py`: escribir una entrada de punta a punta.
    r = subprocess.run([sys.executable, str(AUDIT), "--root", str(d)],
                       input=json.dumps({"command": "aisdd init",
                                         "started_at": "2026-09-03T10:00:00Z"}),
                       capture_output=True, text=True, timeout=60)
    if r.returncode != 0:
        errors.append(f"audit.py falla al ejecutarse: {r.stderr.strip()[-400:]}")

    # 4. `aiba functional-design`: genera el .docx, con plantilla y sin ella.
    #    Necesita python-docx; si no esta, se **dice** y no se finge cobertura.
    df_ejercitado = False
    try:
        import docx                                            # noqa: F401,PLC0415
    except ImportError:
        pass
    else:
        df_ejercitado = True
        manifiesto = {"proyecto": "P", "titulo": "T", "introduccion": "x",
                      "alcance": "x", "narrativa": {"como": "a", "quiero": "b", "para": "c"},
                      "integraciones": "N/A",
                      "validaciones": {"frontal": "N/A", "core": "N/A"},
                      "mensajes": {"frontal": "N/A", "integracion_no_core": "N/A",
                                   "core": "N/A"},
                      "pantallas": "N/A", "especificaciones_tecnicas": "N/A"}
        (d / "m.json").write_text(json.dumps(manifiesto), encoding="utf-8")

        # Plantilla con estilo en espanol y relleno, como la de un cliente.
        import docx as _docx                                   # noqa: PLC0415
        tpl = _docx.Document()
        tpl.add_paragraph("RELLENO DE LA PLANTILLA")
        tpl.styles["Heading 1"].name = "Título 1"
        tpl.save(str(d / "tpl.docx"))

        for etiqueta, extra in (("sin plantilla", []),
                                ("con plantilla", ["--plantilla", str(d / "tpl.docx")])):
            salida = d / f"df-{etiqueta.split()[0]}.docx"
            r = subprocess.run([sys.executable, str(DF), "--manifest", str(d / "m.json"),
                                "--output", str(salida), "--no-install"] + extra,
                               capture_output=True, text=True, timeout=120)
            if r.returncode != 0:
                errors.append(f"gen_df_docx.py falla {etiqueta}: {r.stderr.strip()[-300:]}")
                continue
            doc = _docx.Document(str(salida))
            titulos = [p.text for p in doc.paragraphs
                       if p.style.name in ("Heading 1", "Título 1")]
            if not any(x.startswith("1. ") for x in titulos):
                errors.append(f"gen_df_docx.py {etiqueta}: los apartados no salen "
                              f"numerados ({titulos[:3]})")
            if "con plantilla" in etiqueta:
                if any("RELLENO DE LA PLANTILLA" in p.text for p in doc.paragraphs):
                    errors.append("gen_df_docx.py: el contenido de ejemplo de la "
                                  "plantilla acaba dentro del DF")

if errors:
    print("Scripts que compilan pero no funcionan:", file=sys.stderr)
    for e in errors:
        print(f"  - {e}", file=sys.stderr)
    sys.exit(1)
# Decir que se ejercito y que no. Un "correcto" que oculta una parte sin
# ejecutar es el mismo fallo que esta comprobacion existe para cazar.
print("Humo de scripts correcto: compute_kpis (json y md) y audit.py se ejecutan "
      "de punta a punta sobre un proyecto minimo"
      + (", y gen_df_docx con plantilla y sin ella." if df_ejercitado
         else ". AVISO: gen_df_docx **no** se ha ejercitado (falta python-docx)."))
