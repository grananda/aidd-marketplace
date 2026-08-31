#!/usr/bin/env python3
"""Marca corporativa de los entregables de AIBA, en Word y en Excel.

Vive a nivel de plugin, junto a `stamp_doc.py`, y por la misma razon: la usan
skills distintos y una copia por skill se queda atras sin que nada falle.

El principio, comun a los dos formatos: **la marca se aplica a los estilos, no
al contenido**. En Word, a `Heading 1/2/3` y `Title`; en Excel, a estilos de
celda con nombre. Asi, cambiar de identidad visual mas tarde es cambiar el
estilo, y no repasar el documento parrafo a parrafo o celda a celda.

Sin `branding`, todo lo de aqui es inocuo: los documentos salen neutros, que es
el default de los skills que lo usan.
"""

from __future__ import annotations

import sys
from pathlib import Path

# Paleta neutra. No es "sin color": es una gris deliberada, legible al imprimir
# en blanco y negro y que no compite con ninguna identidad que se aplique luego.
NEUTRO = {
    "cabecera_fondo": "D9D9D9",
    "cabecera_texto": "000000",
    "acento": "404040",
    "banda": "F2F2F2",
}


def _hex(valor: str | None) -> str | None:
    """`#1f3864` -> `1F3864`. Devuelve None si no es un color de 6 digitos."""
    if not valor:
        return None
    v = str(valor).lstrip("#").strip().upper()
    if len(v) == 8:  # ARGB de Excel: se queda el RGB
        v = v[2:]
    if len(v) != 6 or any(c not in "0123456789ABCDEF" for c in v):
        return None
    return v


def normalizar(branding: dict | None) -> dict:
    """Deja el bloque en una forma de la que los generadores puedan fiarse.

    Un color mal escrito se descarta en vez de propagarse: un `RGBColor` que
    revienta a mitad de la generacion deja el fichero a medias, y el color es
    lo menos importante del documento.
    """
    b = dict(branding or {})
    principal = _hex(b.get("color_principal"))
    secundario = _hex(b.get("color_secundario")) or principal
    logo = b.get("logo")
    if logo and not Path(logo).is_file():
        sys.stderr.write(f"Aviso: el logo '{logo}' no existe; se genera sin el.\n")
        logo = None
    return {
        "color_principal": principal,
        "color_secundario": secundario,
        "logo": logo,
        "texto_cabecera": b.get("texto_cabecera") or "",
        "texto_pie": b.get("texto_pie") or "",
        "activa": bool(principal or logo or b.get("texto_cabecera") or b.get("texto_pie")),
    }


def paleta(branding: dict) -> dict:
    """Los cuatro colores que usan las tablas y las hojas, con marca o sin ella."""
    b = normalizar(branding) if "activa" not in (branding or {}) else branding
    principal = b.get("color_principal")
    secundario = b.get("color_secundario") or principal
    if not principal:
        return dict(NEUTRO)
    return {
        "cabecera_fondo": principal,
        "cabecera_texto": "FFFFFF",
        "acento": secundario or principal,
        "banda": "F2F2F2",
    }


# --- Word --------------------------------------------------------------------

def aplicar_estilos_word(doc, branding: dict) -> None:
    """Tine los estilos de titulo. Sin color, no toca nada."""
    from docx.shared import RGBColor

    principal = branding.get("color_principal")
    secundario = branding.get("color_secundario") or principal
    for nombre, color in (("Title", principal), ("Heading 1", principal),
                          ("Heading 2", secundario), ("Heading 3", secundario)):
        if not color:
            continue
        try:
            doc.styles[nombre].font.color.rgb = RGBColor.from_string(color)
        except (KeyError, ValueError):
            pass  # el estilo no existe en la plantilla base; no es motivo de fallo


def cabecera_pie_word(doc, branding: dict, cabecera: str, pie: str) -> None:
    """Cabecera y pie como campos de Word, no como texto fijo.

    El numero de pagina va como campo `PAGE` a proposito: escrito a mano seria
    correcto solo en la pagina uno.
    """
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Cm

    seccion = doc.sections[0]

    p = seccion.header.paragraphs[0]
    p.text = branding.get("texto_cabecera") or cabecera
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    logo = branding.get("logo")
    if logo:
        try:
            p.insert_paragraph_before().add_run().add_picture(logo, height=Cm(1.2))
        except Exception:  # noqa: BLE001 - un logo ilegible no tumba el documento
            sys.stderr.write(f"Aviso: no se pudo insertar el logo '{logo}'; se omite.\n")

    pf = seccion.footer.paragraphs[0]
    pf.text = (branding.get("texto_pie") or pie) + "  ·  Página "
    pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = pf.add_run()
    for tipo, texto in (("begin", None), (None, "PAGE"), ("end", None)):
        el = OxmlElement("w:fldChar") if tipo else OxmlElement("w:instrText")
        if tipo:
            el.set(qn("w:fldCharType"), tipo)
        else:
            el.text = texto
        run._r.append(el)


def sombrear_celda_word(celda, color_hex: str) -> None:
    """Sombreado de celda de tabla. python-docx no lo expone; hay que bajar a OOXML."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    celda._tc.get_or_add_tcPr().append(shd)


# --- Excel -------------------------------------------------------------------

def estilos_excel(branding: dict) -> dict:
    """Fuentes y rellenos listos para openpyxl, derivados de la paleta."""
    from openpyxl.styles import Alignment, Border, Font, PatternFill, Side

    p = paleta(branding)
    linea = Side(style="thin", color="BFBFBF")
    return {
        "cabecera_fuente": Font(bold=True, color=p["cabecera_texto"], size=10),
        "cabecera_relleno": PatternFill("solid", fgColor=p["cabecera_fondo"]),
        "grupo_fuente": Font(bold=True, color=p["acento"], size=11),
        "grupo_relleno": PatternFill("solid", fgColor=p["banda"]),
        "titulo_fuente": Font(bold=True, color=p["acento"], size=14),
        "celda_fuente": Font(size=10),
        "borde": Border(left=linea, right=linea, top=linea, bottom=linea),
        "ajuste": Alignment(vertical="top", wrap_text=True),
        "centro": Alignment(horizontal="center", vertical="center", wrap_text=True),
        "paleta": p,
    }


def logo_excel(ws, branding: dict, ancla: str = "A1") -> bool:
    """Inserta el logo en una hoja. Devuelve si lo consiguio.

    Depende de Pillow, que openpyxl no arrastra. Si falta, el fichero sale sin
    logo y con un aviso: un entregable sin logo sirve, uno a medias no.
    """
    logo = branding.get("logo")
    if not logo:
        return False
    try:
        from openpyxl.drawing.image import Image

        img = Image(logo)
        img.height = min(img.height, 60)
        img.width = int(img.width * (60 / max(img.height, 1))) if img.height else img.width
        ws.add_image(img, ancla)
        return True
    except Exception:  # noqa: BLE001 - falta Pillow, o el fichero no es una imagen
        sys.stderr.write(
            f"Aviso: no se pudo insertar el logo '{logo}' en el Excel "
            f"(¿falta Pillow?); se genera sin el.\n")
        return False
