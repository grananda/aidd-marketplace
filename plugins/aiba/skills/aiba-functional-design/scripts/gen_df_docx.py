#!/usr/bin/env python3
"""aiba-functional-design · gen_df_docx.py — Documento de Diseno Funcional (DF).

Renderiza un manifiesto JSON a un `.docx` con la estructura acordada. Existe un
unico generador a proposito: el valor del skill es que todos los DF de un
proyecto salgan iguales, y eso no se sostiene si cada invocacion arma el
documento a su manera.

El diseno es **generico**: estilos nativos de Word (Titulo 1/2/3, un estilo de
tabla con nombre, cabecera y pie), sin logotipos ni colores corporativos salvo
que el manifiesto traiga una seccion `branding`. Asi, aplicar despues una
identidad visual es cambiar los estilos, no repasar el documento parrafo a
parrafo.

Uso:
    python3 gen_df_docx.py --manifest df.json --output "docs/df/HU-01 - ....docx"
    python3 gen_df_docx.py --schema          # imprime el esquema del manifiesto

Solo depende de `python-docx`, que instala sobre la marcha si falta (misma
estrategia que `gen_hu_plan_xlsx.py`: la invocacion de python es la puerta de
permisos y el pip hereda esa aprobacion). Se puede desactivar con `--no-install`
o la variable de entorno `AIBA_DF_NO_INSTALL`.
"""

from __future__ import annotations

import argparse
import json
import os
import subprocess
import sys
from datetime import date
from pathlib import Path

# La marca vive a nivel de plugin: la comparten este skill y `aiba-test-plan`.
sys.path.insert(0, str(Path(__file__).resolve().parents[3] / "scripts"))

import branding as marca  # noqa: E402 - despues de fijar sys.path

# --- Dependencia -------------------------------------------------------------

def _ensure_docx(allow_install: bool) -> None:
    """Importa python-docx, instalandolo al vuelo si falta."""
    try:
        import docx  # noqa: F401
        return
    except ImportError:
        pass

    if not allow_install or os.environ.get("AIBA_DF_NO_INSTALL"):
        sys.stderr.write(
            "Falta 'python-docx' y la instalacion automatica esta desactivada.\n"
            "Instalalo con:  pip install python-docx\n"
        )
        sys.exit(2)

    sys.stderr.write("Aviso: 'python-docx' no esta instalado; instalandolo automaticamente...\n")
    for cmd in (
        [sys.executable, "-m", "pip", "install", "--quiet", "python-docx"],
        [sys.executable, "-m", "pip", "install", "--quiet", "--user", "python-docx"],
    ):
        try:
            subprocess.check_call(cmd)
        except Exception:  # noqa: BLE001 - se prueba la siguiente estrategia
            continue
        try:
            import docx  # noqa: F401
            sys.stderr.write("OK: 'python-docx' instalado correctamente.\n")
            return
        except ImportError:
            continue
    sys.stderr.write("No se pudo instalar 'python-docx'. Instalalo a mano y reintenta.\n")
    sys.exit(2)


SCHEMA = """\
Esquema del manifiesto (JSON). Todo lo que falte se omite o sale como pendiente.

{
  "proyecto":  "SUPLEMENTOS",             # nombre corto; va en portada y cabecera
  "hu_id":     "HU-03",
  "titulo":    "Busqueda de Poliza",
  "version":   "1.0",
  "autor":     "Nombre Apellido",
  "fecha":     "2026-08-27",              # opcional; por defecto, hoy

  "control_versiones":   [{"fecha": "...", "version": "1.0",
                           "autor": "...", "cambio": "Version inicial"}],
  "control_aprobaciones":[{"responsable": "", "cargo": "",
                           "departamento": "", "fecha": "", "version": ""}],

  "introduccion": "parrafo o lista de parrafos",
  "alcance":      "parrafo o lista de parrafos",

  "narrativa": {"como": "...", "quiero": "...", "para": "..."},

  "campos": {                              # tabla Filtros/Campos
    "columnas": ["Nombre","Editable","Oblig","Tipo","Comentario"],
    "filas": [["Ramo","Si","Si","Lista desplegable","..."]]
  },
  "integraciones": "texto o lista",
  "validaciones": {"frontal": "texto o lista", "core": "texto o lista"},
  "mensajes":     {"frontal": "...", "integracion_no_core": "...", "core": "..."},
  "pantallas":    "texto o lista",
  "imagenes":     ["docs/prototipos/hu-03.png"],

  "criterios_aceptacion": {"contexto": "...", "escenarios": ["Escenario X: ..."]},
  "especificaciones_tecnicas": "texto o lista",
  "puntos_abiertos": [{"id":"PA-01","descripcion":"...","estado":"Abierto",
                       "responsable":"","estimada":"","resolucion":""}],

  "secciones_adicionales": [{"titulo":"Glosario","contenido":"texto o lista"}],

  "branding": {                            # opcional; sin el, documento neutro
    "color_principal":  "1F3864",
    "color_secundario": "2E74B5",
    "logo":             "ruta/al/logo.png",
    "texto_cabecera":   "...",
    "texto_pie":        "..."
  }
}
"""

# El SKILL.md evita tildes por compatibilidad entre plataformas de agentes, pero
# el DOCUMENTO GENERADO lo lee y lo firma un cliente: ahi el espanol va con sus
# tildes. La regla aplica al contenido de salida, no al codigo ni a las instrucciones.
PENDIENTE = "[PENDIENTE: sin información en la documentación de origen]"


# --- Utilidades de contenido -------------------------------------------------

def as_blocks(value) -> list[str]:
    """Normaliza texto suelto o lista a una lista de parrafos no vacios."""
    if value is None:
        return []
    if isinstance(value, str):
        return [b.strip() for b in value.split("\n") if b.strip()]
    if isinstance(value, (list, tuple)):
        out: list[str] = []
        for v in value:
            out.extend(as_blocks(v))
        return out
    return [str(value)]


def write_blocks(doc, value, vacio: str = PENDIENTE, est=None) -> None:
    """Escribe parrafos; las lineas que empiecen por '- ' salen como vinetas."""
    blocks = as_blocks(value)
    if not blocks:
        doc.add_paragraph(vacio)
        return
    for b in blocks:
        if b.startswith(("- ", "* ")):
            doc.add_paragraph(b[2:].strip(),
                              style=est("List Bullet") if est else "List Bullet")
        else:
            doc.add_paragraph(b)


def add_table(doc, columnas: list[str], filas: list[list[str]], accent: str | None) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    t = doc.add_table(rows=1, cols=len(columnas))
    t.style = "Table Grid"
    for i, c in enumerate(columnas):
        celda = t.rows[0].cells[i]
        celda.text = ""
        run = celda.paragraphs[0].add_run(str(c))
        run.bold = True
        celda.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        if accent:
            marca.sombrear_celda_word(celda, accent)
    for fila in filas or []:
        celdas = t.add_row().cells
        for i, v in enumerate(fila[: len(columnas)]):
            celdas[i].text = "" if v is None else str(v)
    doc.add_paragraph()


def add_toc(doc) -> None:
    """Indice como campo TOC: Word lo actualiza solo, escrito a mano se desfasa."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    p = doc.add_paragraph()
    run = p.add_run()
    for tipo, texto in (("begin", None), (None, r'TOC \o "1-3" \h \z \u'), ("separate", None)):
        el = OxmlElement("w:fldChar") if tipo else OxmlElement("w:instrText")
        if tipo:
            el.set(qn("w:fldCharType"), tipo)
        else:
            el.set(qn("xml:space"), "preserve")
            el.text = texto
        run._r.append(el)
    aviso = OxmlElement("w:t")
    aviso.text = "Actualiza el índice en Word: clic derecho > Actualizar campos."
    run._r.append(aviso)
    fin = OxmlElement("w:fldChar")
    fin.set(qn("w:fldCharType"), "end")
    run._r.append(fin)


# --- Documento ---------------------------------------------------------------

# Nombres de estilo por idioma. Word los traduce, asi que una plantilla en
# espanol trae `Titulo 1` donde el script escribe `Heading 1`. Escribir contra un
# nombre que la plantilla no tiene revienta o, peor, deja el parrafo sin formato
# y el documento parece correcto hasta que alguien lo abre.
EQUIVALENTES = {
    "Title":       ["Title", "Título", "Titulo"],
    "Subtitle":    ["Subtitle", "Subtítulo", "Subtitulo"],
    "Heading 1":   ["Heading 1", "Título 1", "Titulo 1"],
    "Heading 2":   ["Heading 2", "Título 2", "Titulo 2"],
    "Heading 3":   ["Heading 3", "Título 3", "Titulo 3"],
    "List Bullet": ["List Bullet", "Lista con viñetas", "Lista con vinetas",
                    "List Paragraph", "Párrafo de lista"],
}


class Estilos:
    """Traduce los nombres logicos a los que existen de verdad en el documento.

    Lo que no puede es fallar en silencio: cada estilo ausente se **reporta** y
    su parrafo se escribe sin estilo, en vez de reventar a mitad del documento.
    """

    def __init__(self, doc) -> None:
        disponibles = {s.name for s in doc.styles}
        self.mapa = {k: next((c for c in v if c in disponibles), None)
                     for k, v in EQUIVALENTES.items()}
        self.faltan = sorted(k for k, v in self.mapa.items() if v is None)

    def __call__(self, logico: str):
        return self.mapa.get(logico, logico)


def limpiar_cuerpo(doc) -> None:
    """Vacia la plantilla conservando estilos, cabecera, pie y formato de pagina.

    Una plantilla suele traer texto de ejemplo; sin quitarlo el DF sale detras.
    Se conserva el `sectPr` final, que es donde viven margenes, tamano y
    orientacion: borrarlo devolveria el documento a los valores por defecto y se
    perderia justo lo que aporta la plantilla.
    """
    cuerpo = doc.element.body
    for hijo in list(cuerpo):
        if not hijo.tag.endswith("}sectPr"):
            cuerpo.remove(hijo)


def titulador(doc, est, numerar: bool = True):
    """Titulos numerados y con el estilo que exista.

    Word no numera los estilos `Heading` por si solo --hace falta una lista
    multinivel vinculada, que en python-docx es XML a mano--. El numero literal
    encaja aqui porque el documento se **regenera**, nunca se edita a mano: no
    hay renumeracion que mantener. Con una plantilla cuyos estilos ya numeren se
    desactiva desde el manifiesto, o saldria `1. 1. Introduccion`.
    """
    estado = {1: 0, 2: 0, 3: 0}

    def escribe(texto: str, nivel: int) -> None:
        if numerar and nivel in estado:
            estado[nivel] += 1
            for menor in range(nivel + 1, 4):
                estado[menor] = 0
            prefijo = ".".join(str(estado[n]) for n in range(1, nivel + 1))
            texto = f"{prefijo}. {texto}" if nivel == 1 else f"{prefijo} {texto}"
        doc.add_paragraph(texto, style=est(f"Heading {nivel}"))

    return escribe


def build(m: dict, salida: Path, plantilla: Path | None = None) -> dict:
    from docx import Document

    branding = marca.normalizar(m.get("branding"))
    accent = (branding.get("color_secundario") or branding.get("color_principal") or "D9D9D9")
    accent = accent.lstrip("#").upper()

    proyecto = m.get("proyecto", "")
    titulo = m.get("titulo") or m.get("hu_id") or "Documento de Diseno Funcional"
    version = str(m.get("version", "1.0"))
    hoy = m.get("fecha") or date.today().isoformat()

    avisos: list[str] = []
    if plantilla:
        if not plantilla.is_file():
            raise SystemExit(f"No existe la plantilla '{plantilla}'.")
        doc = Document(str(plantilla))
        limpiar_cuerpo(doc)
    else:
        doc = Document()
    est = Estilos(doc)
    if est.faltan:
        avisos.append("estilos que la plantilla no trae (esas partes salen sin formato): "
                      + ", ".join(est.faltan))
    h = titulador(doc, est, numerar=bool(m.get("numerar_apartados", True)))
    if branding:
        marca.aplicar_estilos_word(doc, branding)
    marca.cabecera_pie_word(doc, branding,
                            cabecera=f"{proyecto} · {titulo}".strip(" ·"),
                            pie=f"Versión {version}")

    # Portada
    if proyecto:
        doc.add_paragraph(proyecto, style=est("Title"))
    doc.add_paragraph(titulo, style=est("Title") if not proyecto else est("Subtitle"))
    doc.add_paragraph(f"Documento de Diseño Funcional · Versión {version} · {hoy}")
    doc.add_paragraph()

    doc.add_paragraph("Control de Versiones", style=est("Heading 2"))
    add_table(doc, ["Fecha", "Versión", "Autor", "Descripción del cambio"],
              [[c.get("fecha", ""), c.get("version", ""), c.get("autor", ""), c.get("cambio", "")]
               for c in m.get("control_versiones") or
               [{"fecha": hoy, "version": version, "autor": m.get("autor", ""),
                 "cambio": "Version inicial"}]], accent)

    doc.add_paragraph("Control de Aprobaciones", style=est("Heading 2"))
    aprob = m.get("control_aprobaciones") or [{}, {}, {}]
    add_table(doc, ["Responsable", "Cargo", "Departamento", "Fecha", "Versión del documento"],
              [[a.get("responsable", ""), a.get("cargo", ""), a.get("departamento", ""),
                a.get("fecha", ""), a.get("version", "")] for a in aprob], accent)

    doc.add_paragraph("Índice", style=est("Heading 2"))
    add_toc(doc)
    doc.add_page_break()

    # 1. Introduccion
    h("Introducción", 1)
    write_blocks(doc, m.get("introduccion"), est=est)
    h("Alcance", 2)
    write_blocks(doc, m.get("alcance"), est=est)

    # 2. La historia
    h(titulo, 1)
    nar = m.get("narrativa") or {}
    if any(nar.values()):
        for etiqueta, clave in (("COMO", "como"), ("QUIERO", "quiero"), ("PARA", "para")):
            p = doc.add_paragraph()
            p.add_run(f"{etiqueta} ").bold = True
            p.add_run(nar.get(clave, ""))
    else:
        doc.add_paragraph(PENDIENTE)

    h("Filtros/Campos", 2)
    campos = m.get("campos") or {}
    columnas = campos.get("columnas") or ["Nombre", "Editable", "Oblig", "Tipo", "Comentario"]
    if campos.get("filas"):
        add_table(doc, columnas, campos["filas"], accent)
    else:
        doc.add_paragraph("N/A")

    h("Integraciones otros aplicativos", 2)
    write_blocks(doc, m.get("integraciones"), est=est)

    h("Validaciones / Reglas / Acciones", 2)
    val = m.get("validaciones") or {}
    h("Específicas del Frontal", 3)
    write_blocks(doc, val.get("frontal"), est=est)
    h("Específicas del Core", 3)
    write_blocks(doc, val.get("core"), est=est)

    h("Mensajes y avisos", 2)
    msg = m.get("mensajes") or {}
    h("Específicos del Frontal", 3)
    write_blocks(doc, msg.get("frontal"), est=est)
    h("Específicos de Integración no Core", 3)
    write_blocks(doc, msg.get("integracion_no_core"), est=est)
    h("Específicos del Core", 3)
    write_blocks(doc, msg.get("core"), est=est)

    h("Pantallas y Prototipo", 2)
    write_blocks(doc, m.get("pantallas"), est=est)
    from docx.shared import Cm
    for img in m.get("imagenes") or []:
        if Path(img).is_file():
            try:
                doc.add_picture(img, width=Cm(15))
            except Exception:  # noqa: BLE001
                doc.add_paragraph(f"[No se pudo insertar la imagen: {img}]")
        else:
            doc.add_paragraph(f"[Imagen no encontrada: {img}]")

    # 3-5
    h("Criterios de aceptación", 1)
    ca = m.get("criterios_aceptacion") or {}
    if ca.get("contexto"):
        write_blocks(doc, ca["contexto"], est=est)
    escenarios = ca.get("escenarios") or []
    if escenarios:
        for e in escenarios:
            doc.add_paragraph(str(e), style=est("List Bullet"))
    elif not ca.get("contexto"):
        doc.add_paragraph(PENDIENTE)

    h("Especificaciones Técnicas", 1)
    write_blocks(doc, m.get("especificaciones_tecnicas"), est=est)

    h("Puntos abiertos", 1)
    pa = m.get("puntos_abiertos") or []
    add_table(doc, ["ID", "Descripción", "Estado", "Responsable", "F. Estimada", "F. Resolución"],
              [[p.get("id", ""), p.get("descripcion", ""), p.get("estado", "Abierto"),
                p.get("responsable", ""), p.get("estimada", ""), p.get("resolucion", "")]
               for p in pa], accent)

    for extra in m.get("secciones_adicionales") or []:
        h(extra.get("titulo", "Anexo"), 1)
        write_blocks(doc, extra.get("contenido"), est=est)

    salida.parent.mkdir(parents=True, exist_ok=True)
    doc.save(salida)
    return {"output": str(salida), "puntos_abiertos": len(pa),
            "secciones_adicionales": len(m.get("secciones_adicionales") or []),
            "plantilla": str(plantilla) if plantilla else None,
            # Los estilos que la plantilla no traia. Sin reportarlos, el
            # documento sale con partes sin formato y nadie se entera hasta
            # abrirlo.
            "avisos": avisos}


def extraer(ruta: Path) -> dict:
    """El texto y las tablas de un DF ya generado, para que otro skill los lea.

    Un `.docx` es un zip de XML: nadie lo lee de un vistazo. `aiba-test-plan`
    dice que el DF es su mejor fuente --sus validaciones y sus mensajes ya son
    casos de prueba casi literales-- y sin esto esa frase seria decorativa.
    """
    import docx

    d = docx.Document(ruta)
    parrafos = {p._p: p for p in d.paragraphs}
    tablas = {tb._tbl: tb for tb in d.tables}

    secciones: list[dict] = []
    actual = {"titulo": "(portada)", "nivel": 0, "parrafos": [], "tablas": []}
    for bloque in d.element.body.iterchildren():
        p = parrafos.get(bloque)
        if p is not None:
            if not p.text.strip():
                continue
            if p.style.name.startswith("Heading"):
                secciones.append(actual)
                sufijo = p.style.name.split()[-1]
                actual = {"titulo": p.text.strip(),
                          "nivel": int(sufijo) if sufijo.isdigit() else 1,
                          "parrafos": [], "tablas": []}
            else:
                actual["parrafos"].append(p.text.strip())
            continue
        tb = tablas.get(bloque)
        if tb is not None:
            actual["tablas"].append([[c.text.strip() for c in fila.cells]
                                     for fila in tb.rows])
    secciones.append(actual)
    return {"fichero": str(ruta),
            "secciones": [s for s in secciones if s["parrafos"] or s["tablas"]]}


def main() -> int:
    ap = argparse.ArgumentParser(description="Genera el DF en Word de una historia de usuario.")
    ap.add_argument("--manifest", help="JSON con el contenido del DF; sin el, stdin")
    ap.add_argument("--output", help="ruta del .docx de salida")
    ap.add_argument("--schema", action="store_true", help="imprime el esquema del manifiesto y sale")
    ap.add_argument("--extraer", metavar="RUTA",
                    help="vuelca a JSON el texto y las tablas de un DF ya generado, "
                         "para que otro skill pueda leerlo, y sale")
    ap.add_argument("--plantilla", default=None,
                    help="plantilla .docx/.dotx del cliente: el documento hereda sus "
                         "estilos, cabecera, pie y formato de pagina")
    ap.add_argument("--no-install", action="store_true", help="no instalar python-docx al vuelo")
    args = ap.parse_args()

    if args.schema:
        print(SCHEMA)
        return 0
    if args.extraer:
        _ensure_docx(not args.no_install)
        ruta = Path(args.extraer)
        if not ruta.is_file():
            sys.stderr.write(f"No existe el DF '{ruta}'.\n")
            return 2
        print(json.dumps(extraer(ruta), ensure_ascii=False, indent=2))
        return 0
    if not args.output:
        ap.error("--output es obligatorio (o usa --schema)")

    raw = Path(args.manifest).read_text(encoding="utf-8") if args.manifest else sys.stdin.read()
    try:
        m = json.loads(raw)
    except json.JSONDecodeError as exc:
        sys.stderr.write(f"Manifiesto JSON invalido: {exc}\n")
        return 2

    _ensure_docx(allow_install=not args.no_install)
    print(json.dumps(build(m, Path(args.output),
                          Path(args.plantilla) if args.plantilla else None),
                     ensure_ascii=False))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
