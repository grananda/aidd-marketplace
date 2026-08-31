#!/usr/bin/env python3
"""aiba-test-plan · gen_evidence_docx.py — Documento de Evidencias de una HU.

Renderiza el mismo manifiesto que el plan a un `.docx` con un bloque por caso de
prueba, listo para que quien ejecuta pegue la captura bajo "Resultado actual".

La plantilla de referencia es un **combinador de correspondencia** (`«Código»`,
`«Pasos»`, `«Próximo registro»`) que se rellena en Word desde el Excel. Aqui el
documento sale ya expandido: los datos los tenemos, y obligar a combinar a mano
es un paso manual sin nada a cambio.

Uso:
    python3 gen_evidence_docx.py --manifest plan.json --output "docs/pruebas/Evidencias - HU-06.docx"
    python3 gen_evidence_docx.py --schema

El esquema del manifiesto es el de `gen_test_plan_xlsx.py --schema`: **un solo
manifiesto para los dos entregables**, para que no puedan describir cosas
distintas. Solo depende de `python-docx`, que instala al vuelo si falta.
"""

from __future__ import annotations

import argparse
import json
import os
import subprocess
import sys
from datetime import date
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parents[3] / "scripts"))
sys.path.insert(0, str(Path(__file__).resolve().parent))


def _ensure_docx(allow_install: bool) -> None:
    try:
        import docx  # noqa: F401
        return
    except ImportError:
        pass

    if not allow_install or os.environ.get("AIBA_PP_NO_INSTALL"):
        sys.stderr.write(
            "Falta 'python-docx' y la instalacion automatica esta desactivada.\n"
            "Instalalo con:  pip install python-docx\n")
        sys.exit(2)

    sys.stderr.write("Aviso: 'python-docx' no esta instalado; instalandolo automaticamente...\n")
    for cmd in ([sys.executable, "-m", "pip", "install", "--quiet", "python-docx"],
                [sys.executable, "-m", "pip", "install", "--quiet", "--user", "python-docx"]):
        try:
            subprocess.check_call(cmd)
        except Exception:  # noqa: BLE001
            continue
        try:
            import docx  # noqa: F401
            sys.stderr.write("OK: 'python-docx' instalado correctamente.\n")
            return
        except ImportError:
            continue
    sys.stderr.write("No se pudo instalar 'python-docx'. Instalalo a mano y reintenta.\n")
    sys.exit(2)


# --- Utilidades --------------------------------------------------------------

def parrafos(valor) -> list[str]:
    if valor is None:
        return []
    if isinstance(valor, (list, tuple)):
        out: list[str] = []
        for v in valor:
            out.extend(parrafos(v))
        return out
    return [b.strip() for b in str(valor).split("\n") if b.strip()]


def escribir(doc, valor, vacio: str = "") -> None:
    bloques = parrafos(valor)
    if not bloques:
        if vacio:
            doc.add_paragraph(vacio)
        return
    for b in bloques:
        if b.startswith(("- ", "* ")):
            doc.add_paragraph(b[2:].strip(), style="List Bullet")
        else:
            doc.add_paragraph(b)


def etiqueta(doc, texto: str) -> None:
    """`Descripción:` en negrita, como en la plantilla de referencia."""
    p = doc.add_paragraph()
    p.add_run(texto).bold = True


def tabla(doc, filas: list[list[str]], anchos: list[int] | None, marca, br: dict,
          cabecera: bool = True, combinar_cabecera: bool = False) -> None:
    from docx.shared import Cm

    t = doc.add_table(rows=0, cols=len(filas[0]))
    t.style = "Table Grid"
    paleta = marca.paleta(br)
    for n, fila in enumerate(filas):
        celdas = t.add_row().cells
        for i, v in enumerate(fila[: len(celdas)]):
            celdas[i].text = ""
            run = celdas[i].paragraphs[0].add_run("" if v is None else str(v))
            if cabecera and n == 0:
                run.bold = True
                marca.sombrear_celda_word(celdas[i], paleta["cabecera_fondo"])
    if combinar_cabecera and len(t.rows) > 1:
        # El rotulo de la ficha ocupa la fila entera. Sin combinar quedan tres
        # celdas vacias a su derecha, que es como se ve en la plantilla cuando
        # la combinacion de correspondencia no llego a rellenarlas.
        t.rows[0].cells[0].merge(t.rows[0].cells[-1])
    if anchos:
        for fila in t.rows:
            for i, ancho in enumerate(anchos[: len(fila.cells)]):
                fila.cells[i].width = Cm(ancho)
    doc.add_paragraph()


def indice(doc) -> None:
    """Indice como campo TOC: Word lo actualiza solo; a mano se desfasa."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    p = doc.add_paragraph()
    run = p.add_run()
    for tipo, texto in (("begin", None), (None, r'TOC \o "1-3" \h \z \u'), ("separate", None)):
        el = OxmlElement("w:fldChar") if tipo else OxmlElement("w:instrText")
        if tipo:
            el.set(qn("w:fldCharType"), tipo)
        else:
            el.set(qn("xml:space"), "preserve")
            el.text = texto
        run._r.append(el)
    aviso = OxmlElement("w:t")
    aviso.text = "Actualiza el índice en Word: clic derecho > Actualizar campos."
    run._r.append(aviso)
    fin = OxmlElement("w:fldChar")
    fin.set(qn("w:fldCharType"), "end")
    run._r.append(fin)


# --- Documento ---------------------------------------------------------------

def build(m: dict, salida: Path) -> dict:
    import branding as marca
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    from gen_test_plan_xlsx import generar_codigos

    br = marca.normalizar(m.get("branding"))
    hoy = m.get("fecha") or date.today().isoformat()
    proyecto = m.get("proyecto", "")
    hu = m.get("hu_id", "")
    titulo = m.get("titulo", "")
    version = str(m.get("version", "1.0"))
    casos = list(m.get("casos") or [])
    codigos = generar_codigos(casos)
    nombres_cu = {cu.get("id"): cu.get("nombre", "") for cu in m.get("casos_uso") or []}

    doc = Document()
    if br.get("activa"):
        marca.aplicar_estilos_word(doc, br)
    marca.cabecera_pie_word(
        doc, br,
        cabecera=f"{proyecto} · Evidencias {hu}".strip(" ·"),
        pie=f"Versión {version}")

    # Portada
    doc.add_paragraph("Documento de Evidencias de Pruebas", style="Title")
    sub = doc.add_paragraph(f"{proyecto} — {hu} {titulo}".strip(" —"))
    sub.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if m.get("referencia"):
        doc.add_paragraph(f"Referencia: {m['referencia']}")
    doc.add_paragraph(f"Versión {version} · {hoy}")
    doc.add_paragraph()
    indice(doc)
    doc.add_page_break()

    # Datos generales
    doc.add_heading("Datos Generales", level=1)
    tabla(doc, [
        ["Campo", "Valor"],
        ["Proyecto", proyecto],
        ["Historia de usuario", f"{hu} {titulo}".strip()],
        ["Referencia externa", m.get("referencia", "")],
        ["Entorno", m.get("entorno", "")],
        ["Total de casos de prueba", str(len(casos))],
        ["Fecha final", ""],
        ["Resultado final", ""],
    ], [5.5, 10.5], marca, br)
    doc.add_paragraph(
        "Las filas «Fecha final» y «Resultado final» las completa quien ejecuta, "
        "al cerrar la campaña de pruebas.")

    # Pruebas funcionales
    doc.add_heading("Pruebas Funcionales", level=1)

    doc.add_heading("Catálogo de evidencias", level=2)
    filas = [["Caso de Uso", "Nombre", "Descripción y escenarios", "Id. Requisito", "Validación"]]
    for cu in m.get("casos_uso") or []:
        req = cu.get("requisitos") or []
        filas.append([
            cu.get("id", ""), cu.get("nombre", ""), " ".join(parrafos(cu.get("descripcion"))),
            ", ".join(str(r) for r in req) if isinstance(req, (list, tuple)) else str(req),
            "",
        ])
    tabla(doc, filas, [2.2, 3.6, 6.0, 2.2, 2.0], marca, br)

    doc.add_heading("Evidencias", level=2)
    doc.add_paragraph(
        f"A continuación se recogen los {len(casos)} casos de prueba de la historia "
        f"{hu}. Bajo «Resultado actual» de cada caso se adjunta la evidencia de su "
        f"ejecución.")

    for caso, codigo in zip(casos, codigos):
        cu = caso.get("caso_uso", "")
        nombre = caso.get("nombre", "")
        doc.add_heading(f"{cu} / {codigo} / {nombre}".strip(" /"), level=3)
        tabla(doc, [
            ["DOCUMENTACIÓN DE PRUEBAS FUNCIONALES", "", "", ""],
            ["Id caso de prueba:", codigo, "Nombre caso de uso:",
             f"{cu} - {nombres_cu.get(cu, '')}".strip(" -")],
            ["Descripción del caso de uso:", " ".join(parrafos(caso.get("descripcion"))),
             "Nombre caso de prueba:", nombre],
            ["Autor:", m.get("autor", ""), "Fecha:", hoy],
            ["Criticidad:", caso.get("criticidad", ""), "Tipo de ejecución:",
             caso.get("ejecucion", "manual")],
        ], [4.0, 4.2, 4.0, 4.2], marca, br, combinar_cabecera=True)

        for rotulo, valor in (("Descripción:", caso.get("descripcion")),
                              ("Precondiciones:", caso.get("precondiciones")),
                              ("Pasos:", caso.get("pasos")),
                              ("Resultado esperado:", caso.get("resultado_esperado"))):
            etiqueta(doc, rotulo)
            escribir(doc, valor, vacio="—")

        etiqueta(doc, "Resultado actual:")
        doc.add_paragraph()  # el hueco de la captura
        doc.add_paragraph()

    for p in doc.paragraphs:
        for run in p.runs:
            if run.font.size is None and p.style.name == "Normal":
                run.font.size = Pt(10)

    salida.parent.mkdir(parents=True, exist_ok=True)
    doc.save(salida)
    return {
        "salida": str(salida),
        "casos": len(casos),
        "casos_uso": len(m.get("casos_uso") or []),
        "marca": bool(br.get("activa")),
    }


def main() -> None:
    ap = argparse.ArgumentParser(
        description="Genera el Documento de Evidencias (.docx) de una HU.")
    ap.add_argument("--manifest", help="el mismo JSON que el plan; sin el, stdin")
    ap.add_argument("--output", help="ruta del .docx de salida")
    ap.add_argument("--schema", action="store_true",
                    help="recuerda donde esta el esquema y sale")
    ap.add_argument("--no-install", action="store_true",
                    help="no instalar python-docx al vuelo")
    args = ap.parse_args()

    if args.schema:
        print("El manifiesto es el mismo que el del plan. Consultalo con:\n"
              "  python3 gen_test_plan_xlsx.py --schema")
        return
    if not args.output:
        ap.error("--output es obligatorio (o usa --schema)")

    _ensure_docx(not args.no_install)

    crudo = (Path(args.manifest).read_text(encoding="utf-8") if args.manifest
             else sys.stdin.read())
    try:
        m = json.loads(crudo)
    except json.JSONDecodeError as e:
        sys.stderr.write(f"El manifiesto no es JSON valido: {e}\n")
        sys.exit(2)

    if not m.get("casos"):
        sys.stderr.write(
            "El manifiesto no trae ningun caso de prueba: no hay evidencias que "
            "recoger. Revisa los criterios de aceptacion de la historia.\n")
        sys.exit(2)

    print(json.dumps(build(m, Path(args.output)), ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
